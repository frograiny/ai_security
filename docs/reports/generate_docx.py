import os, re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_final_report():
    md_file = r'd:\AI\ai_security\docs\reports\BaoCao_CuoiKy_HeThongBaoMatWeb_BiLSTM.md'
    docx_file = r'd:\AI\ai_security\docs\reports\BaoCao_CuoiKy_HeThongBaoMatWeb_BiLSTM.docx'
    
    if not os.path.exists(md_file):
        print(f"Error: {md_file} not found.")
        return

    doc = Document()
    
    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        if line.startswith('# '):
            h = doc.add_heading(line.replace('# ', ''), level=1)
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=2)
        elif line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=3)
        elif line.startswith('---'):
            doc.add_page_break()
        elif line.startswith('* '):
            doc.add_paragraph(line.replace('* ', ''), style='List Bullet')
        elif re.match(r'^\d+\.\s', line):
            doc.add_paragraph(line, style='List Number')
        else:
            doc.add_paragraph(line)

    doc.save(docx_file)
    print(f"Successfully created {docx_file}")

if __name__ == "__main__":
    create_final_report()
